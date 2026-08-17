'''Ingestion seam: directory walk, file classification, frontmatter + section
parsing, structured record building (markdown and V3 packs), PDF text/anchor
extraction, ID-mention scanning, and the top-level scan_and_parse entrypoint.'''
import hashlib
import logging
import os
import re
import uuid
from datetime import date, datetime, timezone
from pathlib import Path

import frontmatter
import yaml

from cofr.domain import DOMAIN_TYPES, TYPE_TO_COLLECTION, from_dict, validate_and_normalize
from cofr.packs import pack_load


HARDCODED_EXCLUDED_DIRS = frozenset({'.git', '.cofr', '__pycache__', '.venv', 'node_modules', 'artifacts'})
TEXT_EXTENSIONS = frozenset({
    '.md', '.markdown', '.txt', '.json', '.yaml', '.yml', '.csv', '.tsv',
    '.py', '.r', '.jl', '.sh', '.ts', '.js', '.go', '.rs', '.c', '.cpp', '.h',
    '.html', '.css', '.toml', '.ini', '.cfg', '.rst', '.org', '.tex', '.bib',
})
IMAGE_EXTENSIONS = frozenset({'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tif', '.tiff', '.webp'})
OFFICE_EXTENSIONS = frozenset({'.docx', '.xlsx'})
CONTENT_EXTRACTED_EXTENSIONS = frozenset({'.pdf'}) | IMAGE_EXTENSIONS | OFFICE_EXTENSIONS
LARGE_FILE_THRESHOLD = 100 * 1024 * 1024

_PACK_CANONICAL_FILENAMES = frozenset({'claims.yaml', 'decisions.yaml', 'questions.yaml', 'risks.yaml', 'experiments.yaml'})
_PACK_FILENAME_TO_TYPE = {
    'claims.yaml': 'claim',
    'decisions.yaml': 'decision',
    'questions.yaml': 'question',
    'risks.yaml': 'risk',
    'experiments.yaml': 'experiment',
}

_HEADING_TRAIL_PUNCT = '?.!:'
_VALID_ID_RE = re.compile(r'^[A-Za-z0-9_-]+$')
_FENCE_RE = re.compile(r'^(?P<fence>`{3,}|~{3,})')
_BULLET_RE = re.compile(r'^\s*-\s+(.*)$')
_ID_BOUNDARY_RE_CACHE = {}

PARSE_MODES = ('prose', 'reference_list', 'polarity_list', 'key_value_dict', 'string_list')

SECTION_MAP = {
    'claim': {
        'title': ('title', 'prose'),
        'statement': ('statement', 'prose'),
        'main_support': ('main_support', 'prose'),
        'main_weakness': ('main_weakness', 'prose'),
        'what_would_change_my_mind': ('what_would_change_my_mind', 'prose'),
    },
    'evidence': {
        'summary': ('summary', 'prose'),
        'affects_claims': ('claim_links', 'polarity_list'),
    },
    'experiment': {
        'name': ('name', 'prose'),
        'intent': ('intent', 'prose'),
        'config_reference': ('config_reference', 'prose'),
        'config': ('config_reference', 'prose'),
        'result_summary': ('result_summary', 'prose'),
        'result': ('result_summary', 'prose'),
        'key_metrics': ('key_metrics', 'key_value_dict'),
        'metrics': ('key_metrics', 'key_value_dict'),
        'implications': ('implications', 'prose'),
        'affected_claims': ('affected_claim_ids', 'reference_list'),
        'affects_claims': ('affected_claim_ids', 'reference_list'),
        'follow_on_questions': ('follow_on_questions', 'string_list'),
    },
    'decision': {
        'title': ('title', 'prose'),
        'decision_statement': ('decision_statement', 'prose'),
        'statement': ('decision_statement', 'prose'),
        'rationale': ('rationale', 'prose'),
        'based_on_evidence': ('based_on_evidence_ids', 'reference_list'),
        'based_on': ('based_on_evidence_ids', 'reference_list'),
        'depends_on_claims': ('depends_on_claim_ids', 'reference_list'),
        'depends_on': ('depends_on_claim_ids', 'reference_list'),
        'reopen_conditions': ('reopen_conditions', 'prose'),
        'reopen_when': ('reopen_conditions', 'prose'),
    },
    'question': {
        'question': ('question', 'prose'),
        'blocking_impact': ('blocking_impact', 'prose'),
        'impact': ('blocking_impact', 'prose'),
        'related_claims': ('related_claim_ids', 'reference_list'),
        'proposed_resolution': ('proposed_resolution', 'prose'),
        'minimum_test': ('minimum_test', 'prose'),
    },
    'risk': {
        'statement': ('statement', 'prose'),
        'related_claims': ('related_claim_ids', 'reference_list'),
        'related_decisions': ('related_decision_ids', 'reference_list'),
        'recommended_resolution': ('recommended_resolution', 'prose'),
        'resolution': ('recommended_resolution', 'prose'),
    },
}

_TYPE_AWARE_STALE_WARNINGS = {
    'claim': 'stale is system-managed; to deactivate a claim, set status: retired',
    'evidence': 'stale is system-managed; to deactivate evidence, set status: deprecated',
    'decision': 'stale is system-managed; to deactivate a decision, set status: deprecated',
    'experiment': 'stale is system-managed; to deactivate an experiment, set status: deprecated',
    'question': 'stale is system-managed; to deactivate a question, set status: resolved or deprioritized',
    'risk': 'stale is system-managed; to deactivate a risk, set status: mitigated or resolved',
}


def _stale_warning_for_type(type_str):
    return _TYPE_AWARE_STALE_WARNINGS.get(type_str, 'stale is system-managed; user-authored stale is ignored')


def _normalize_heading(raw):
    text = raw.strip()
    text = text.rstrip(_HEADING_TRAIL_PUNCT).strip()
    text = ' '.join(text.split())
    return text.lower().replace(' ', '_')


def _stringify_dates(metadata):
    out = {}
    for k, v in metadata.items():
        if isinstance(v, (date, datetime)):
            out[k] = v.isoformat()
        else:
            out[k] = v
    return out


def _slug_from(s):
    out = ''.join(ch if (ch.isascii() and ch.isalnum()) or ch == '_' else '_' for ch in str(s).strip().lower())
    while '__' in out:
        out = out.replace('__', '_')
    return out.strip('_')


def split_frontmatter(text):
    '''Parse YAML frontmatter from text.

    Returns (metadata_dict, body_str). When no frontmatter exists, metadata is {}
    and body is the unchanged input. When YAML fails to parse, metadata is
    {'_yaml_error': <message>} and body is the unchanged input so the caller can
    still inspect or surface it. YAML date/datetime scalars are converted to
    ISO-8601 strings so they round-trip cleanly through JSON state.
    '''
    normalized = text.replace('\r\n', '\n').replace('\r', '\n')
    try:
        post = frontmatter.loads(normalized)
    except yaml.YAMLError as exc:
        return {'_yaml_error': str(exc)}, normalized
    metadata = _stringify_dates(dict(post.metadata)) if post.metadata else {}
    return metadata, post.content


def parse_sections(body):
    '''Split markdown body into {normalized_heading: content_str}.

    Sections are separated by H2 lines (`## Heading`) that appear outside fenced
    code blocks. Headings are normalized to snake_case for lookup. Content before
    the first H2 is ignored. Section content has leading and trailing blank lines
    stripped but preserves internal whitespace.
    '''
    sections = {}
    current_heading = None
    current_lines = []
    in_fence = False
    fence_char = ''

    def flush():
        if current_heading is None:
            return
        content = '\n'.join(current_lines).strip('\n').strip()
        sections[current_heading] = content

    for line in body.splitlines():
        fence_match = _FENCE_RE.match(line.lstrip())
        if fence_match:
            marker = fence_match.group('fence')[0]
            if not in_fence:
                in_fence = True
                fence_char = marker
            elif marker == fence_char:
                in_fence = False
                fence_char = ''
            if current_heading is not None:
                current_lines.append(line)
            continue
        if not in_fence and line.startswith('## '):
            flush()
            current_heading = _normalize_heading(line[3:])
            current_lines = []
            continue
        if current_heading is not None:
            current_lines.append(line)

    flush()
    return sections


def _bullet_items(content):
    items = []
    for line in content.splitlines():
        m = _BULLET_RE.match(line) or re.match(r'^\s*-(\S.*)$', line)
        if m:
            items.append(m.group(1).strip())
    return items


def parse_field_value(mode, content):
    '''Parse a section's content into the typed value for a field.

    Returns (value, [warning]). Returns (None, [warning]) when the content
    cannot be coerced into the target shape; the caller is responsible for
    deciding what to do with the raw section content (typically: store in
    extra_sections so nothing is lost).
    '''
    if mode == 'prose':
        return content.strip('\n').strip(), []

    if mode not in PARSE_MODES:
        return None, [f'unknown parse mode {mode!r}']

    if not content.strip():
        return None, [f'expected bullet list for parse mode {mode!r}; section is empty']

    items = _bullet_items(content)
    if not items:
        return None, [f'expected bullet list for parse mode {mode!r}; no bullets found']

    if mode == 'reference_list':
        return [item.strip() for item in items if item.strip()], []

    if mode == 'string_list':
        return items, []

    if mode == 'polarity_list':
        out = []
        warnings = []
        for item in items:
            if ':' in item:
                head, _, tail = item.rpartition(':')
                pol = tail.strip() or 'supports'
                if pol not in ('supports', 'opposes'):
                    warnings.append(f'invalid polarity {pol!r} on link to {head.strip()!r}; defaulting to supports (valid: supports, opposes)')
                    pol = 'supports'
                out.append({'claim_id': head.strip(), 'polarity': pol})
            else:
                out.append({'claim_id': item.strip(), 'polarity': 'supports'})
        return out, warnings

    if mode == 'key_value_dict':
        out = {}
        for item in items:
            if ':' not in item:
                return None, [f'malformed key-value bullet (no colon): {item!r}']
            label, _, value = item.partition(':')
            out[label.strip()] = value.strip()
        if not out:
            return None, ['key-value section had no parseable bullets']
        return out, []

    return None, [f'unhandled parse mode {mode!r}']


def assign_id(metadata, parsed_path, seen_this_refresh, type_str):
    '''Resolve the object ID for a record being parsed.

    Returns (id, collision_flag, warning_or_None, id_was_generated). The 4th
    element is True only when the function minted a fresh UUID (no explicit id
    in the source). The in-refresh map catches duplicate explicit IDs within
    this refresh; state-level duplicate handling lives in the merge step.
    '''
    explicit_present = 'id' in metadata
    explicit = metadata.get('id')
    if explicit_present:
        if not isinstance(explicit, str) or not _VALID_ID_RE.match(explicit):
            return explicit, True, f'invalid id {explicit!r} in {parsed_path!r}: ids must match [A-Za-z0-9_-]+ (no path separators, dots, colons, or whitespace)', False
        earlier_path = seen_this_refresh.get(explicit)
        if earlier_path is not None and earlier_path != parsed_path:
            return explicit, True, f'duplicate id {explicit!r}: earlier file this refresh was {earlier_path!r}, also declared in {parsed_path!r}', False
        return explicit, False, None, False
    return f'{type_str}_{uuid.uuid4().hex}', False, None, True


def _is_text(path):
    try:
        with open(path, 'rb') as fh:
            chunk = fh.read(8192)
        chunk.decode('utf-8')
        return True
    except (UnicodeDecodeError, OSError):
        return False


def _detect_frontmatter(path):
    '''Inspect a file's frontmatter.

    Returns (type_str | None, yaml_error | None). `type_str` is the raw `type:`
    value when frontmatter parses cleanly; otherwise None. `yaml_error` is the
    YAML parse error message when frontmatter is present but malformed.
    '''
    try:
        text = Path(path).read_text(encoding='utf-8')
    except (UnicodeDecodeError, OSError):
        return None, None
    if not text.lstrip().startswith('---'):
        return None, None
    metadata, _ = split_frontmatter(text)
    if '_yaml_error' in metadata:
        return None, metadata['_yaml_error']
    type_str = metadata.get('type')
    return (type_str if type_str else None), None


def _is_pack_path(path, project_path=None):
    '''True iff path is a canonical V3 pack location.'''
    p = Path(path)
    if project_path is not None:
        try:
            rel = p.resolve().relative_to(Path(project_path).resolve())
        except ValueError:
            return False
        parts = rel.parts
        if len(parts) == 1:
            return parts[0] in _PACK_CANONICAL_FILENAMES
        if len(parts) == 2 and parts[0] == 'evidences' and rel.suffix in ('.yaml', '.yml'):
            return True
        return False
    name = p.name
    if name in _PACK_CANONICAL_FILENAMES:
        return True
    return p.suffix in ('.yaml', '.yml') and p.parent.name == 'evidences'


def classify_file(path, size, project_path=None):
    '''Classify one file and (if structured) return its frontmatter type.

    Returns (classification, frontmatter_type | None, yaml_error | None).
    Classifications: structured | structured_pack | unstructured | content_extracted |
    non_text | large_file. Files with malformed YAML frontmatter classify as
    `unstructured` (the parser cannot route them) but `yaml_error` is non-None.
    Files at canonical V3 pack paths (claims.yaml, evidences/*.yaml, etc.) classify
    as `structured_pack`.
    '''
    path = Path(path)
    if size > LARGE_FILE_THRESHOLD:
        return 'large_file', None, None
    ext = path.suffix.lower()
    if ext in CONTENT_EXTRACTED_EXTENSIONS:
        return 'content_extracted', None, None
    if ext in ('.yaml', '.yml') and _is_pack_path(path, project_path):
        return 'structured_pack', None, None
    if ext in TEXT_EXTENSIONS or _is_text(path):
        fm_type, yaml_error = _detect_frontmatter(path)
        if yaml_error:
            return 'unstructured', None, yaml_error
        return ('structured', fm_type, None) if fm_type else ('unstructured', None, None)
    return 'non_text', None, None


def _load_gitignore_patterns(project_path):
    gi = Path(project_path) / '.gitignore'
    if not gi.is_file():
        return []
    patterns = []
    for raw in gi.read_text(encoding='utf-8').splitlines():
        line = raw.strip()
        if not line or line.startswith('#'):
            continue
        patterns.append(line)
    return patterns


def _gitignore_glob_to_regex(pattern):
    parts = []
    i = 0
    while i < len(pattern):
        ch = pattern[i]
        if ch == '*':
            if i + 1 < len(pattern) and pattern[i + 1] == '*':
                i += 2
                if i < len(pattern) and pattern[i] == '/':
                    i += 1
                    parts.append('(?:.*/)?')
                else:
                    parts.append('.*')
                continue
            parts.append('[^/]*')
        elif ch == '?':
            parts.append('[^/]')
        else:
            parts.append(re.escape(ch))
        i += 1
    return ''.join(parts)


def _single_gitignore_pattern_matches(rel_path, pattern):
    pat = pattern.rstrip('/')
    if not pat:
        return False
    directory_pattern = pattern.endswith('/')
    anchored = pat.startswith('/')
    if anchored:
        pat = pat[1:]
    if not pat:
        return False

    has_slash = '/' in pat
    body = _gitignore_glob_to_regex(pat)
    if has_slash or anchored:
        regex = f'^{body}'
    else:
        regex = f'(^|.*/){body}'
    if directory_pattern:
        regex += '(/.*)?$'
    else:
        regex += '$'
    return re.match(regex, rel_path) is not None


def _gitignore_matches(rel_path, patterns):
    matched = False
    for pattern in patterns:
        negated = pattern.startswith('!')
        pat = pattern[1:] if negated else pattern
        if not pat:
            continue
        if _single_gitignore_pattern_matches(rel_path, pat):
            matched = not negated
    return matched


def _page_count_from_text(text):
    '''Count form-feed markers in markitdown-extracted text.'''
    if not text:
        return None
    ff_count = text.count('\f')
    if ff_count > 0:
        return ff_count + 1
    return None


def _page_count_from_pdf_bytes(abs_path):
    '''Fallback: count /Type /Page object headers in raw PDF bytes
    (excluding the /Pages tree parent).'''
    try:
        with open(abs_path, 'rb') as fh:
            data = fh.read()
    except OSError:
        return None
    matches = re.findall(rb'/Type\s*/Page(?![s/A-Za-z0-9])', data)
    if matches:
        return len(matches)
    return None


def extract_pdf_text(path, return_error=False):
    '''Extract text from a PDF via markitdown.

    Returns (text | None, page_count | None) by default. With return_error=True,
    returns (text | None, page_count | None, error_message | None).

    Page-count fallback chain per plan:
      1. Markitdown-extracted text form-feed markers
      2. /Type /Page regex on raw PDF bytes
      3. None
    Page counts are heuristic and intended only for source-anchor validation.
    '''
    def failed(error=None):
        if return_error:
            return None, None, error
        return None, None

    try:
        with open(path, 'rb') as fh:
            magic = fh.read(5)
    except OSError as exc:
        return failed(f'{type(exc).__name__}: {exc}')
    if magic != b'%PDF-':
        return failed('not a PDF file: missing %PDF- header')
    try:
        from markitdown import MarkItDown
        md = MarkItDown()
        pdfminer_logger = logging.getLogger('pdfminer.pdfdocument')
        prior_level = pdfminer_logger.level
        try:
            pdfminer_logger.setLevel(logging.ERROR)
            result = md.convert(str(path))
        finally:
            pdfminer_logger.setLevel(prior_level)
        text = result.text_content
        page_count = _page_count_from_text(text)
        if page_count is None:
            page_count = _page_count_from_pdf_bytes(path)
        if return_error:
            return text, page_count, None
        return text, page_count
    except Exception as exc:
        return failed(f'{type(exc).__name__}: {exc}')


def extract_image_metadata(path):
    '''Return deterministic image metadata, without OCR or interpretation.'''
    try:
        from PIL import ExifTags, Image
        with Image.open(path) as image:
            width, height = image.size
            image_format = image.format or Path(path).suffix.lstrip('.').upper()
            raw_exif = image.getexif()
            image_exif = {}
            for tag_id, value in raw_exif.items():
                tag_name = str(ExifTags.TAGS.get(tag_id, tag_id))
                if value is None or isinstance(value, (str, int, float, bool)):
                    safe_value = value
                elif isinstance(value, bytes):
                    safe_value = value.hex()
                elif isinstance(value, (list, tuple)):
                    safe_value = [item if isinstance(item, (str, int, float, bool)) or item is None else str(item) for item in value]
                else:
                    safe_value = str(value)
                image_exif[tag_name] = safe_value
        return {
            'image_dimensions': [width, height],
            'image_format': image_format,
            'image_exif': dict(sorted(image_exif.items())),
        }, None
    except Exception as exc:
        return None, f'{type(exc).__name__}: {exc}'


def extract_docx_text(path):
    '''Mechanically extract paragraph and table-cell text from DOCX.'''
    try:
        from docx import Document
        document = Document(str(path))
        chunks = [p.text for p in document.paragraphs if p.text]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text for cell in row.cells]
                if any(values):
                    chunks.append('\t'.join(values))
        return '\n'.join(chunks), None
    except Exception as exc:
        return None, f'{type(exc).__name__}: {exc}'


def extract_xlsx_text(path):
    '''Mechanically extract displayed cell values and sheet names from XLSX.'''
    workbook = None
    try:
        from openpyxl import load_workbook
        workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
        sheet_names = list(workbook.sheetnames)
        chunks = []
        for sheet in workbook.worksheets:
            chunks.append(f'## {sheet.title}')
            for row in sheet.iter_rows(values_only=True):
                values = ['' if value is None else str(value) for value in row]
                if any(values):
                    chunks.append('\t'.join(values))
        return '\n'.join(chunks), sheet_names, None
    except Exception as exc:
        return None, None, f'{type(exc).__name__}: {exc}'
    finally:
        if workbook is not None:
            workbook.close()


def validate_anchors_against_index(parsed_records, new_index, project_path):
    '''Pass 2 anchor validation. Returns list of warnings.

    Owns: source_path sandbox shape, symlink-escape, index-miss, per-anchor
    extension match, page-range check.

    Invalid source_path/source_anchors are dropped from the parsed record while
    preserving the evidence record itself.
    '''
    warnings = []
    project_path = Path(project_path).resolve()

    def preserve_user_fields(obj, **fields_to_preserve):
        preserved = dict(getattr(obj, '_preserved_user_fields', {}) or {})
        for key, value in fields_to_preserve.items():
            if value not in ('', None) and value != [] and value != {}:
                preserved[key] = value
        obj._preserved_user_fields = preserved

    for obj in parsed_records:
        if not hasattr(obj, 'source_path'):
            continue
        src_path = getattr(obj, 'source_path', '') or ''
        anchors = list(getattr(obj, 'source_anchors', []) or [])
        if not src_path and not anchors:
            continue
        cleared_reason = None
        if anchors and not src_path:
            warnings.append(f'{obj.parsed_from}: source_anchors require source_path; dropping anchors')
            preserve_user_fields(obj, source_anchors=anchors)
            obj.source_anchors = []
            continue
        if src_path:
            if Path(src_path).is_absolute():
                cleared_reason = f'source_path {src_path!r} is absolute; must be project-relative'
            elif '..' in Path(src_path).parts:
                cleared_reason = f'source_path {src_path!r} contains `..`; must not escape project root'
            else:
                try:
                    resolved = (project_path / src_path).resolve(strict=False)
                    resolved.relative_to(project_path)
                except ValueError:
                    cleared_reason = f'source_path {src_path!r} resolves outside project root (symlink escape)'
            if not cleared_reason and src_path not in new_index:
                cleared_reason = f'source_path {src_path!r} not found in index'
        if cleared_reason:
            warnings.append(f'{obj.parsed_from}: {cleared_reason}')
            preserve_user_fields(obj, source_path=src_path, source_anchors=anchors)
            obj.source_path = ''
            obj.source_anchors = []
            continue
        if src_path:
            ext = Path(src_path).suffix.lower()
            page_count = new_index.get(src_path, {}).get('page_count')
            kept = []
            for anchor in anchors:
                if not isinstance(anchor, dict):
                    warnings.append(f'{obj.parsed_from}: anchor must be a mapping; got {type(anchor).__name__}; dropping')
                    continue
                has_page = 'page' in anchor
                has_section = 'section' in anchor
                if has_page and has_section:
                    warnings.append(f'{obj.parsed_from}: anchor has both page and section keys; dropping')
                    continue
                if not has_page and not has_section:
                    warnings.append(f'{obj.parsed_from}: anchor has neither page nor section keys; dropping')
                    continue
                if has_page:
                    if ext != '.pdf':
                        warnings.append(f'{obj.parsed_from}: PDF-shape anchor on non-PDF source; dropping')
                        continue
                    page = anchor.get('page')
                    if not isinstance(page, int):
                        warnings.append(f'{obj.parsed_from}: anchor page must be int; got {type(page).__name__}; dropping')
                        continue
                    if page < 1:
                        warnings.append(f'{obj.parsed_from}: anchor page must be >= 1; dropping')
                        continue
                    if page_count is not None and page > page_count:
                        warnings.append(f'{obj.parsed_from}: anchor page {page} exceeds page_count {page_count}; dropping')
                        continue
                else:
                    if ext == '.pdf':
                        warnings.append(f'{obj.parsed_from}: text-shape anchor on PDF source; dropping')
                        continue
                    if ext not in {'.md', '.markdown', '.txt'}:
                        warnings.append(f'{obj.parsed_from}: text-shape anchor requires markdown or text source; dropping')
                        continue
                    section = anchor.get('section')
                    if not isinstance(section, str) or not section.strip():
                        warnings.append(f'{obj.parsed_from}: anchor section must be a non-empty string; dropping')
                        continue
                    if 'line' in anchor:
                        line = anchor.get('line')
                        if not isinstance(line, int):
                            warnings.append(f'{obj.parsed_from}: anchor line must be int; got {type(line).__name__}; dropping')
                            continue
                        if line < 1:
                            warnings.append(f'{obj.parsed_from}: anchor line must be >= 1; dropping')
                            continue
                kept.append(anchor)
            if kept != anchors:
                preserve_user_fields(obj, source_anchors=anchors)
            obj.source_anchors = kept
    return warnings


def _build_existing_id_to_path(state):
    out = {}
    for key in ('claims', 'evidence', 'experiments', 'decisions', 'open_questions', 'risks'):
        for item in state.get(key, []):
            obj_id = item.get('id')
            parsed_from = item.get('parsed_from')
            if obj_id and parsed_from:
                out[obj_id] = parsed_from
    return out


def scan_for_id_mentions(text, all_ids, min_length=6):
    '''Return the IDs from `all_ids` that appear in `text`.

    IDs shorter than min_length are skipped to avoid false positives. Match
    requires that the ID be bounded on both sides by characters NOT in the
    ID grammar [A-Za-z0-9_-] (so `claim_a` does not match inside
    `claim_a_extended`). Order is deterministic (matches order of all_ids);
    duplicates within text count once.
    '''
    found = []
    for obj_id in all_ids:
        if not isinstance(obj_id, str) or len(obj_id) < min_length:
            continue
        pattern = _ID_BOUNDARY_RE_CACHE.get(obj_id)
        if pattern is None:
            pattern = re.compile(r'(?<![A-Za-z0-9_\-])' + re.escape(obj_id) + r'(?![A-Za-z0-9_\-])')
            _ID_BOUNDARY_RE_CACHE[obj_id] = pattern
        if pattern.search(text):
            found.append(obj_id)
    return found


def inspect_structured_markdown_ids(path):
    '''Parse only frontmatter of a markdown file. Returns (id_or_None, type_or_None).

    Used by scan_existing_ids -- does NOT call assign_id, so no UUIDs are minted.
    '''
    try:
        text = Path(path).read_text(encoding='utf-8')
    except (UnicodeDecodeError, OSError):
        return None, None
    if not text.lstrip().startswith('---'):
        return None, None
    metadata, _ = split_frontmatter(text)
    if '_yaml_error' in metadata:
        return None, None
    return metadata.get('id'), metadata.get('type')


def scan_existing_ids(project_path, state, exclude_patterns=None):
    '''Detect id collisions across state, packs, and structured markdown.

    Returns (id_map, malformed_packs, warnings).
    id_map[id] is a list of occurrence dicts {record_type, in_state, state_parsed_from, disk_locations}.
    Each disk_location is {path, kind, record_index?}.
    exclude_patterns (from .cofr/config.yaml) is forwarded to walk_project so
    config-excluded markdown files don't participate in collision detection.
    '''
    project_path = Path(project_path)
    id_map = {}
    malformed_packs = []
    warnings = []

    state_id_to_path = {}
    state_id_to_type = {}
    for type_str, collection_key in TYPE_TO_COLLECTION.items():
        for item in state.get(collection_key, []):
            rid = item.get('id')
            if rid:
                state_id_to_path[rid] = item.get('parsed_from', '')
                state_id_to_type[rid] = type_str

    disk_occurrences = {}

    pack_files = []
    for pack_name in ('claims.yaml', 'decisions.yaml', 'questions.yaml', 'risks.yaml', 'experiments.yaml'):
        p = project_path / pack_name
        if p.is_file():
            pack_files.append((p, _PACK_FILENAME_TO_TYPE[pack_name]))
    ev_dir = project_path / 'evidences'
    if ev_dir.is_dir():
        for p in sorted([*ev_dir.glob('*.yaml'), *ev_dir.glob('*.yml')]):
            pack_files.append((p, 'evidence'))

    for pack_path, expected_type in pack_files:
        try:
            records, pack_warnings = pack_load(pack_path, expected_type=expected_type, return_warnings=True)
        except Exception as exc:
            malformed_packs.append((str(pack_path.relative_to(project_path)), str(exc)))
            continue
        warnings.extend(pack_warnings)
        try:
            loaded_raw = yaml.safe_load(pack_path.read_text(encoding='utf-8')) or []
        except Exception:
            loaded_raw = records
        records_for_occurrences = loaded_raw if isinstance(loaded_raw, list) else records
        for idx, raw in enumerate(records_for_occurrences):
            if not isinstance(raw, dict):
                continue
            rid = raw.get('id')
            if not rid:
                continue
            rtype = raw.get('type') or expected_type
            if expected_type and rtype != expected_type:
                continue
            disk_occurrences.setdefault(rid, []).append({
                'path': str(pack_path.relative_to(project_path)),
                'kind': 'pack',
                'record_index': idx,
                'record_type': rtype,
            })

    for rel_path, abs_path, size in walk_project(project_path, exclude_patterns=exclude_patterns):
        if abs_path.suffix.lower() != '.md':
            continue
        rid, rtype = inspect_structured_markdown_ids(abs_path)
        if not rtype:
            continue
        if rid:
            disk_occurrences.setdefault(rid, []).append({
                'path': rel_path,
                'kind': 'markdown',
                'record_type': rtype,
            })
        else:
            state_rid = None
            for s_rid, s_path in state_id_to_path.items():
                if s_path == rel_path:
                    state_rid = s_rid
                    break
            if state_rid:
                disk_occurrences.setdefault(state_rid, []).append({
                    'path': rel_path,
                    'kind': 'markdown',
                    'record_type': state_id_to_type.get(state_rid, rtype),
                })
            else:
                warnings.append(f'{rel_path}: structured markdown without explicit id: -- id will be assigned on next `cofr refresh`; this file is not yet eligible for collision detection.')

    all_ids = set(state_id_to_path) | set(disk_occurrences)
    for rid in all_ids:
        in_state = rid in state_id_to_path
        disk_locs_raw = disk_occurrences.get(rid, [])
        rtype = state_id_to_type.get(rid) or (disk_locs_raw[0]['record_type'] if disk_locs_raw else None)
        if in_state and disk_locs_raw:
            state_path = state_id_to_path[rid]
            state_pack_path = state_path.split('#', 1)[0] if '#' in state_path else state_path
            other_locs = [
                {'path': loc['path'], 'kind': loc['kind'], 'record_index': loc.get('record_index')}
                for loc in disk_locs_raw if loc['path'] != state_pack_path
            ]
            on_state_pack = [
                {'path': loc['path'], 'kind': loc['kind'], 'record_index': loc.get('record_index')}
                for loc in disk_locs_raw if loc['path'] == state_pack_path
            ]
            if on_state_pack and other_locs:
                id_map[rid] = [
                    {'record_type': rtype, 'in_state': True, 'state_parsed_from': state_id_to_path[rid], 'disk_locations': on_state_pack},
                ] + [
                    {'record_type': loc['kind'] == 'pack' and 'pack' or 'markdown', 'in_state': False, 'state_parsed_from': None, 'disk_locations': [loc]}
                    for loc in other_locs
                ]
            elif on_state_pack:
                id_map[rid] = [{'record_type': rtype, 'in_state': True, 'state_parsed_from': state_id_to_path[rid], 'disk_locations': on_state_pack}]
            elif other_locs:
                id_map[rid] = [{'record_type': rtype, 'in_state': True, 'state_parsed_from': state_id_to_path[rid], 'disk_locations': other_locs}]
        elif in_state:
            id_map[rid] = [{'record_type': rtype, 'in_state': True, 'state_parsed_from': state_id_to_path[rid], 'disk_locations': []}]
        else:
            grouped_locs = {}
            for loc in disk_locs_raw:
                grouped_locs.setdefault((loc['path'], loc['kind'], loc['record_type']), []).append(
                    {'path': loc['path'], 'kind': loc['kind'], 'record_index': loc.get('record_index')}
                )
            for (path, kind, record_type), disk_locations in grouped_locs.items():
                id_map.setdefault(rid, []).append({
                    'record_type': record_type,
                    'in_state': False,
                    'state_parsed_from': None,
                    'disk_locations': disk_locations,
                })

    return id_map, malformed_packs, warnings


def parse_record_dict(type_str, raw, parsed_from, pack_filename_stem=None):
    '''Convert a YAML record dict into a domain object.

    `raw` is the dict loaded from a pack file. `parsed_from` is the synthetic
    `<pack-path>#<id>` location. `pack_filename_stem` is the stem of the pack
    filename for evidence (used to canonicalize source_slug).

    Returns (obj | None, [warning_str], needs_pack_rewrite: bool).
    needs_pack_rewrite is True iff the parse auto-corrected a field on disk
    (currently: source_slug mismatch).
    '''
    if type_str not in DOMAIN_TYPES:
        valid = ', '.join(sorted(DOMAIN_TYPES))
        return None, [f'unknown type {type_str!r} in {parsed_from!r}; valid types: {valid}'], False
    warnings = []
    rec = dict(raw)
    rec_id = rec.get('id')
    if not isinstance(rec_id, str) or not _VALID_ID_RE.match(rec_id):
        return None, [f'invalid id {rec_id!r} in {parsed_from!r}: ids must match [A-Za-z0-9_-]+ (no path separators, dots, colons, or whitespace)'], False
    if 'stale' in rec:
        warning_text = _stale_warning_for_type(type_str)
        warnings.append(f'{parsed_from}: {warning_text}')
        rec.pop('stale', None)
    if type_str == 'claim' and 'source_missing' in rec:
        warnings.append(f'{parsed_from}: source_missing is system-managed; ignoring user-authored value')
        rec.pop('source_missing', None)
    obj = from_dict(type_str, rec)
    obj.parsed_from = parsed_from
    needs_rewrite = False
    _, shape_warnings = validate_and_normalize(obj, user_authored=True)
    for w in shape_warnings:
        warnings.append(f'{parsed_from}: {w}')
    if type_str == 'evidence':
        anchors = list(obj.source_anchors or [])
        kept = []
        shapes_seen = set()
        for a in anchors:
            if not isinstance(a, dict):
                warnings.append(f'{parsed_from}: anchor {a!r} is not a dict; dropping')
                continue
            has_page = 'page' in a
            has_section = 'section' in a
            if has_page and has_section:
                warnings.append(f'{parsed_from}: anchor {a!r} has both page and section; dropping (anchors are PDF-shape OR text-shape, never both)')
                continue
            if not has_page and not has_section:
                warnings.append(f'{parsed_from}: anchor {a!r} has neither page nor section; dropping')
                continue
            shape = 'pdf' if has_page else 'text'
            shapes_seen.add(shape)
            kept.append((shape, a))
        if len(shapes_seen) > 1:
            warnings.append(f'{parsed_from}: source_anchors mix PDF and text shapes; mixed shapes are not allowed (a record has exactly one source). Dropping all anchors.')
            obj.source_anchors = []
        else:
            obj.source_anchors = [a for _shape, a in kept]
        if pack_filename_stem:
            stored_slug = obj.source_slug
            if stored_slug and stored_slug != pack_filename_stem:
                warnings.append(f'{parsed_from}: source_slug {stored_slug!r} differs from pack filename stem {pack_filename_stem!r}; using {pack_filename_stem!r}')
                needs_rewrite = True
            obj.source_slug = pack_filename_stem
    return obj, warnings, needs_rewrite


def parse_structured_record(path, text, seen_this_refresh):
    '''Parse one structured markdown file into a domain object.

    Returns (dataclass_instance | None, [warning_str], id_was_generated: bool).
    Returns (None, [], False) when the file has no frontmatter (not structured
    -- not an error). Returns (None, [warning_str], False) when the file is
    structured but has an issue (malformed YAML, unknown type, ID collision).
    When a section's content fails per-field parsing, raw text is stored in
    `extra_sections` keyed by the normalized heading; the typed field stays at
    its default. `id_was_generated` is True when the id was minted from a uuid
    (no explicit `id:` in source).
    '''
    metadata, body = split_frontmatter(text)
    if '_yaml_error' in metadata:
        return None, [f'malformed YAML frontmatter in {path!r}: {metadata['_yaml_error']}'], False
    type_str = metadata.get('type')
    if not type_str:
        return None, [], False
    if type_str not in DOMAIN_TYPES:
        valid = ', '.join(sorted(DOMAIN_TYPES))
        return None, [f'unknown type {type_str!r} in {path!r}; valid types: {valid}'], False

    sections = parse_sections(body)
    obj_id, collision, id_warning, id_was_generated = assign_id(metadata, path, seen_this_refresh, type_str)
    if collision:
        return None, [id_warning], False

    cls = DOMAIN_TYPES[type_str]
    obj = cls(id=obj_id)
    obj.parsed_from = path

    warnings = []
    fields_from_frontmatter = set()
    unknown_fields = dict(obj._unknown_fields) if getattr(obj, '_unknown_fields', None) else {}
    for meta_key, meta_value in metadata.items():
        if meta_key in ('type', 'id'):
            continue
        if meta_key == 'stale':
            warnings.append(f'{path}: {_stale_warning_for_type(type_str)}')
            continue
        if type_str == 'claim' and meta_key == 'source_missing':
            warnings.append(f'{path}: source_missing is system-managed; ignoring user-authored value')
            continue
        if hasattr(obj, meta_key) and not meta_key.startswith('_'):
            setattr(obj, meta_key, meta_value)
            fields_from_frontmatter.add(meta_key)
        else:
            unknown_fields[meta_key] = meta_value
    if hasattr(obj, '_unknown_fields'):
        obj._unknown_fields = unknown_fields

    section_map = SECTION_MAP.get(type_str, {})
    extra_sections = dict(obj.extra_sections) if obj.extra_sections else {}
    for heading, content in sections.items():
        mapped = section_map.get(heading)
        if mapped is None:
            extra_sections[heading] = content
            continue
        field_name, mode = mapped
        if field_name in fields_from_frontmatter:
            extra_sections[heading] = content
            continue
        value, parse_warnings = parse_field_value(mode, content)
        for w in parse_warnings:
            warnings.append(f'{path}: {heading!r}: {w}')
        if value is None:
            extra_sections[heading] = content
            continue
        setattr(obj, field_name, value)
    obj.extra_sections = extra_sections

    _, enum_warnings = validate_and_normalize(obj, user_authored=True)
    for w in enum_warnings:
        warnings.append(f'{path}: {w}')

    return obj, warnings, id_was_generated


def walk_project(project_path, exclude_patterns=None):
    '''Yield (rel_path, abs_path, size) for every file under project_path.

    Skips hardcoded excluded directories, .gitignore-matched paths, and any
    paths matched by exclude_patterns (loaded from .cofr/config.yaml). Walks
    in sorted order so behavior is deterministic across OS / filesystem layouts.
    exclude_patterns uses the same matching semantics as .gitignore so a user
    pattern like 'data/**' skips data/ and everything beneath it.
    '''
    project_path = Path(project_path)
    patterns = _load_gitignore_patterns(project_path)
    extra = list(exclude_patterns) if exclude_patterns else []
    for root, dirs, files in os.walk(project_path):
        rel_root = os.path.relpath(root, project_path)
        if rel_root == '.':
            rel_root = ''
        dirs[:] = sorted(d for d in dirs if d not in HARDCODED_EXCLUDED_DIRS)
        if patterns or extra:
            kept_dirs = []
            for d in dirs:
                rel_dir = f'{rel_root}/{d}' if rel_root else d
                if patterns and _gitignore_matches(rel_dir, patterns):
                    continue
                if extra and _gitignore_matches(rel_dir, extra):
                    continue
                kept_dirs.append(d)
            dirs[:] = kept_dirs
        for name in sorted(files):
            rel_path = f'{rel_root}/{name}' if rel_root else name
            if patterns and _gitignore_matches(rel_path, patterns):
                continue
            if extra and _gitignore_matches(rel_path, extra):
                continue
            abs_path = Path(root) / name
            try:
                size = abs_path.stat().st_size
            except OSError:
                continue
            yield rel_path, abs_path, size


def _sha256_hex(path):
    h = hashlib.sha256()
    with open(path, 'rb') as fh:
        for chunk in iter(lambda: fh.read(65536), b''):
            h.update(chunk)
    return h.hexdigest()


def _iso_mtime(path):
    return datetime.fromtimestamp(Path(path).stat().st_mtime, tz=timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')


def scan_and_parse(project_path, existing_state, exclude_patterns=None):
    '''Walk the project, classify files, parse structured records, scan for ID mentions.

    Returns (new_index, parsed_records, warnings, packs_parsed_successfully,
    packs_to_rewrite, id_was_generated_by_path).
    exclude_patterns (from .cofr/config.yaml) is forwarded to walk_project so
    config-excluded paths are skipped from indexing and parsing.
    '''
    project_path = Path(project_path)
    existing_id_to_path = _build_existing_id_to_path(existing_state)
    seen_this_refresh = {}
    new_index = {}
    parsed_records = []
    warnings = []
    text_for_scanning = {}
    packs_parsed_successfully = set()
    packs_to_rewrite = set()
    id_was_generated_by_path = {}

    for rel_path, abs_path, size in walk_project(project_path, exclude_patterns=exclude_patterns):
        classification, fm_type, yaml_error = classify_file(abs_path, size, project_path=project_path)
        if yaml_error:
            warnings.append(f'malformed YAML frontmatter in {rel_path!r}: {yaml_error}')
        try:
            content_hash = None if classification in ('non_text', 'large_file') else _sha256_hex(abs_path)
            mtime = _iso_mtime(abs_path)
        except OSError as exc:
            warnings.append(f'failed to stat {rel_path!r}: {exc}')
            continue
        new_index[rel_path] = {
            'mtime': mtime,
            'size': size,
            'content_hash': content_hash,
            'classification': classification,
            'extension': abs_path.suffix.lower(),
            'frontmatter_type': fm_type,
            'id_mentions': [],
            'page_count': None,
            'extracted_text_length': None,
            'image_dimensions': None,
            'image_format': None,
            'image_exif': None,
            'sheet_names': None,
        }
        if classification == 'structured':
            try:
                text = abs_path.read_text(encoding='utf-8')
            except (UnicodeDecodeError, OSError) as exc:
                warnings.append(f'failed to read {rel_path!r}: {exc}')
                continue
            obj, parse_warnings, id_was_generated = parse_structured_record(rel_path, text, seen_this_refresh)
            warnings.extend(parse_warnings)
            if obj is not None:
                parsed_records.append(obj)
                seen_this_refresh[obj.id] = rel_path
                id_was_generated_by_path[rel_path] = id_was_generated
        elif classification == 'structured_pack':
            filename = abs_path.name
            if abs_path.parent.name == 'evidences':
                expected_type = 'evidence'
                pack_stem = abs_path.stem
            else:
                expected_type = _PACK_FILENAME_TO_TYPE.get(filename)
                pack_stem = None
            try:
                records, pack_warnings = pack_load(abs_path, expected_type=expected_type, return_warnings=True)
            except Exception as exc:
                warnings.append(f'{rel_path}: malformed YAML pack: {exc}')
                continue
            for w in pack_warnings:
                warnings.append(w if rel_path in w else f'{rel_path}: {w}')
            packs_parsed_successfully.add(rel_path)
            for raw in records:
                rid = raw.get('id')
                rtype = raw.get('type') or expected_type
                synthetic_parsed_from = f'{rel_path}#{rid}'
                obj, record_warnings, needs_rewrite = parse_record_dict(rtype, raw, synthetic_parsed_from, pack_filename_stem=pack_stem)
                warnings.extend(record_warnings)
                if obj is None:
                    continue
                if needs_rewrite:
                    packs_to_rewrite.add(rel_path)
                earlier = seen_this_refresh.get(rid)
                if earlier is not None and earlier != synthetic_parsed_from:
                    warnings.append(f'duplicate id {rid!r}: earlier this refresh was {earlier!r}, also declared in {synthetic_parsed_from!r}')
                    continue
                parsed_records.append(obj)
                seen_this_refresh[rid] = synthetic_parsed_from
        elif classification == 'content_extracted':
            extension = abs_path.suffix.lower()
            error = None
            if extension == '.pdf':
                text, page_count, error = extract_pdf_text(abs_path, return_error=True)
                if text is not None:
                    text_for_scanning[rel_path] = text
                    new_index[rel_path]['page_count'] = page_count
                    new_index[rel_path]['extracted_text_length'] = len(text)
            elif extension in IMAGE_EXTENSIONS:
                metadata, error = extract_image_metadata(abs_path)
                if metadata is not None:
                    new_index[rel_path].update(metadata)
                    new_index[rel_path]['extracted_text_length'] = 0
            elif extension == '.docx':
                text, error = extract_docx_text(abs_path)
                if text is not None:
                    text_for_scanning[rel_path] = text
                    new_index[rel_path]['extracted_text_length'] = len(text)
            elif extension == '.xlsx':
                text, sheet_names, error = extract_xlsx_text(abs_path)
                if text is not None:
                    text_for_scanning[rel_path] = text
                    new_index[rel_path]['extracted_text_length'] = len(text)
                    new_index[rel_path]['sheet_names'] = sheet_names
            if error is not None:
                new_index[rel_path]['classification'] = 'non_text'
                new_index[rel_path]['content_hash'] = None
                warnings.append(f'{rel_path}: {extension.lstrip(".").upper()} extraction failed ({error}); classified as non_text')
        elif classification == 'unstructured':
            try:
                text_for_scanning[rel_path] = abs_path.read_text(encoding='utf-8')
            except (UnicodeDecodeError, OSError):
                pass

    known_ids = set(seen_this_refresh) | set(existing_id_to_path)
    all_ids = sorted(known_ids)

    for rel_path, text in text_for_scanning.items():
        new_index[rel_path]['id_mentions'] = scan_for_id_mentions(text, all_ids)

    return new_index, parsed_records, warnings, packs_parsed_successfully, packs_to_rewrite, id_was_generated_by_path
